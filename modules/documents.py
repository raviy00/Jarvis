# modules/documents.py - Export replies as Word Documents
import os
from datetime import datetime
from docx import Document

OUTPUT_DIR = "outputs"

def save_as_docx(text, filename=None):
    """Save text content into a .docx file."""
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)
        
    if not filename:
        filename = f"jarvis_output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    path = os.path.join(OUTPUT_DIR, filename)
    
    try:
        doc = Document()
        doc.add_heading('JARVIS - Academic Report', 0)
        
        # Add generated content
        doc.add_paragraph(text)
        
        doc.add_page_break()
        doc.save(path)
        return path
    except Exception as e:
        print(f"❌ Document save error: {e}")
        return None
